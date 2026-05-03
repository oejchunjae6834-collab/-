# -*- coding: utf-8 -*-
"""
동아리모임 회의록 docx 생성 스크립트.
sample 형식(동아리모임_샘플.md)에 맞춰 표/헤딩/리스트로 정리.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\\Users\\경남교육청\\Desktop\\claude-test\\outputs\\동아리모임_20260428.docx"

doc = Document()

# 기본 폰트 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(11)

# 제목
title = doc.add_heading('디적디적 운영진 회의록', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 회의 개요
doc.add_heading('회의 개요', level=1)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
data = [
    ('일시', '2026년 4월 28일 (화) 20:02 ~ 21:18 (75분)'),
    ('장소', '온라인 (구글 미트)'),
    ('참석자', '오은진, 양미선, 이산지나, 고경애'),
    ('작성자', '오은진'),
]
for i, (k, v) in enumerate(data):
    table.cell(i, 0).text = k
    table.cell(i, 1).text = v

# 1. 재정 현황 및 회비·캠프비 논의
doc.add_heading('1. 재정 현황 및 회비·캠프비 논의', level=1)

doc.add_heading('현황', level=2)
for line in [
    '공금 잔액: 약 130만 원 (정산 전 기준)',
    '미정산 항목: 산아 책방 카드 사용분, 오은진 선생님 선물, 이사장님 차비',
    '정산 후 실제 잔액은 약 절반 수준 예상',
    '회원들의 자발적 반찬·간식·쌀 찬조가 운영비 절감에 큰 도움',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('회비 구조 분석', level=2)
for line in [
    '6개월치 회비 + 캠프 1회 = 거의 소진되는 구조',
    '주요 지출: 이사장님 숙박비(주말 요금 적용), 식사 지원비',
    '이사장님 1차 식비는 공금, 2차(카페 등)는 이사장님 제외 N분의 1',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('회비·캠프비 인상 논의', level=2)
for line in [
    '월 회비(2만 원): CMS 추가 부담(소요 조합비 등) 고려해 현행 유지',
    '여름 캠프비(1일 일정): 현행 4만 원 유지 검토',
    '겨울 캠프비(2일 일정): 4만 원 → 5만 원 인상 방향',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('결정사항', level=2)
for line in [
    '월 회비 현행 유지',
    '겨울 캠프비 5만 원 인상 방향 잠정 합의',
    '5월 모임에서 여름 캠프 계획 확정 후 최종 결정',
    '이사장님 문화상품권 지출 증가 → 수량 조율 필요',
]:
    doc.add_paragraph(line, style='List Bullet')

# 2. 5월 산아 책방 모임 (회장 선출 모임)
doc.add_heading('2. 5월 산아 책방 모임 (회장 선출 모임)', level=1)

doc.add_heading('진행 방식', level=2)
for line in [
    '5월: 이사장님 미참석 월 → 정기 모임 대신 산아 책방 브런치 모임으로 진행',
    '차기 회장 선출이 주요 안건',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('예산', level=2)
for line in [
    '브런치 1인 2만 원 공금 지원 (정기 모임 대체 성격)',
    '참석 예상 인원: 운영진 4명 + 이은숙 선생님 외',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('결정사항', level=2)
for line in [
    '공금으로 브런치 진행',
    '단톡방에 회장 선출 모임이라는 점 안내하며 참석 독려 공지 게시',
]:
    doc.add_paragraph(line, style='List Bullet')

# 3. 차기 회장 선출 방식
doc.add_heading('3. 차기 회장 선출 방식', level=1)

doc.add_heading('선출 방식', level=2)
for line in [
    '자원자 부재로 제비뽑기 방식 결정',
    '뽑기 대상: 운영진 4명 + 이은숙, 유광미, 산아 선생님 = 7명 안 (차미란 선생님 포함 시 8명)',
    '제비 7개 준비 예정',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('1인 1역할 분담제 도입', level=2)
for line in [
    '회장: 제비뽑기로 결정',
    '총무: 양미선 선생님 2년 연임 → 3선 불가, 신규 선출 필요',
    '캠프 부장: 부회장 대신 명확한 역할명으로 신설 (회장 부담 분산)',
    '정기 모임 기록: 오은진 선생님 자원 (회장 여부와 무관하게 지속)',
    '출석 체크: 이산지나 선생님 자원 → 카톡 복붙 방식 개선, 별도 프로그램 검토',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('모임 정체성 재확인', level=2)
for line in [
    '"아이들 교육 모임"이 아닌 "어른이 먼저 배우는 모임"',
    '아이 참석 여부와 무관하게 어른의 학습·성장이 중심',
    '신규 회원에게도 어른 공부 모임 참여가 본질이라는 점 안내 필요',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('결정사항', level=2)
for line in [
    '5월 산아 책방 모임에서 제비뽑기로 차기 회장 선출',
    '회장·총무·캠프 부장·기록·출석 체크 등 1인 1역할 업무 분장',
    '각자 맡을 역할 사전 검토 후 5월 모임에서 확정',
]:
    doc.add_paragraph(line, style='List Bullet')

# 4. 휴면 회원 관련 논의
doc.add_heading('4. 휴면 회원 관련 논의', level=1)
for line in [
    '김지혜 선생님: 지난 캠프 이후 회비·캠프비 미납, 연락 두절 상태',
    '양미선 관장님이 카톡·전화로 활동 의사 재확인 시도',
    '유광미 선생님: 본인이 "바빠서 당분간 못 온다"고 의사 표현 → 활동 의지 있는 것으로 판단',
    '예린이(김지혜 선생님 자녀) 버즈 분실물: 발견 시 김지혜 선생님 측에 전달 예정',
]:
    doc.add_paragraph(line, style='List Bullet')

# 5. 기타 사항
doc.add_heading('5. 기타 사항', level=1)
for line in [
    '이사장님 제안: 여름방학 중학생 대상 심도 있는 학습 프로그램(산아 책방 활용) 검토',
    '운영진 부담 완화 필요성 공감 → 즐겁게 지속 가능한 운영 방향 모색',
    '5월 정기 모임에서 여름 캠프 계획 함께 논의 예정',
]:
    doc.add_paragraph(line, style='List Bullet')

# 6. 다음 일정
doc.add_heading('6. 다음 일정', level=1)
table2 = doc.add_table(rows=3, cols=2)
table2.style = 'Light Grid Accent 1'
hdr = table2.rows[0].cells
hdr[0].text = '일정'
hdr[1].text = '주요 안건'
schedule = [
    ('5월 산아 책방 모임', '제비뽑기로 차기 회장 선출, 1인 1역할 업무 분장 확정'),
    ('5월 정기 모임', '여름 캠프 계획 수립, 회비·캠프비 최종 결정'),
]
for i, (k, v) in enumerate(schedule, start=1):
    table2.cell(i, 0).text = k
    table2.cell(i, 1).text = v

doc.save(OUT)
print('saved:', OUT)
